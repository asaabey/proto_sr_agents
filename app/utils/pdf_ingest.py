"""
PDF/Word document ingestion utilities for systematic review manuscripts.

Implemented:
- python-docx for Word document structure parsing
- spaCy NLP for PICO element extraction
- Pattern matching for search strategies and flow diagrams

Future:
- GROBID for PDF structured extraction
- pdfplumber/camelot for table extraction  
- Tesseract for OCR of scanned PDFs
"""

from typing import Dict, List, Optional, Any, Tuple
from pathlib import Path
import re
from collections import defaultdict

# Optional dependencies
try:
    from docx import Document
    from docx.table import Table
    DOCX_AVAILABLE = True
    # Type hints for docx objects
    DocxDocument = Document
except ImportError:
    DOCX_AVAILABLE = False
    # Fallback for type hints when docx not available
    class DocxDocument: pass
    class Table: pass

try:
    import spacy
    NLP_AVAILABLE = True
except ImportError:
    NLP_AVAILABLE = False

from app.models.schemas import Manuscript, PICO, SearchDescriptor, FlowCounts, StudyRecord, OutcomeEffect, ExclusionReason

class TextExtractor:
    """Common text processing utilities for both PDF and Word processors."""
    
    def __init__(self):
        self.nlp = None
        if NLP_AVAILABLE:
            try:
                self.nlp = spacy.load("en_core_web_sm")
            except OSError:
                print("Warning: spaCy model 'en_core_web_sm' not found. Some NLP features disabled.")
    
    def extract_pico_elements(self, text: str) -> Optional[PICO]:
        """Extract PICO elements from manuscript text using pattern matching and NLP."""
        
        # Pattern-based extraction
        pico_data = {
            'framework': 'PICO',
            'population': None,
            'intervention': None, 
            'comparator': None,
            'outcomes': []
        }
        
        # Look for explicit PICO sections
        pico_patterns = {
            'population': r'(?:population|participants?|patients?)[:\s]+([^.]+)',
            'intervention': r'(?:intervention|treatment|therapy)[:\s]+([^.]+)', 
            'comparator': r'(?:comparator?|control|comparison)[:\s]+([^.]+)',
            'outcomes': r'(?:outcomes?|endpoints?)[:\s]+([^.]+)'
        }
        
        for element, pattern in pico_patterns.items():
            matches = re.findall(pattern, text, re.IGNORECASE | re.MULTILINE)
            if matches:
                if element == 'outcomes':
                    # Split outcomes by common delimiters
                    outcomes_text = matches[0]
                    outcomes = [o.strip() for o in re.split(r'[;,:]|\band\b', outcomes_text) if o.strip()]
                    pico_data['outcomes'] = outcomes[:5]  # Limit to first 5
                else:
                    pico_data[element] = matches[0].strip()
        
        # Enhanced extraction using NLP if available
        if self.nlp and any(pico_data.values()):
            pico_data = self._enhance_pico_with_nlp(text, pico_data)
        
        # Only return PICO if we found substantial content
        if sum(1 for v in pico_data.values() if v) >= 2:
            return PICO(**pico_data)
        
        return None
    
    def _enhance_pico_with_nlp(self, text: str, pico_data: Dict[str, Any]) -> Dict[str, Any]:
        """Use NLP to enhance PICO extraction."""
        doc = self.nlp(text[:10000])  # Limit text length for processing
        
        # Extract medical entities
        medical_entities = []
        for ent in doc.ents:
            if ent.label_ in ["PERSON", "ORG", "PRODUCT", "EVENT", "DISEASE"]:
                medical_entities.append(ent.text)
        
        # If population is missing, look for demographic mentions
        if not pico_data['population'] and medical_entities:
            population_candidates = [ent for ent in medical_entities if any(
                keyword in ent.lower() for keyword in ['patient', 'adult', 'child', 'participant']
            )]
            if population_candidates:
                pico_data['population'] = population_candidates[0]
        
        return pico_data
    
    def parse_search_strategies(self, text: str) -> List[SearchDescriptor]:
        """Parse search strategy sections to extract database searches.""" 
        strategies = []
        
        # Common database patterns
        db_patterns = {
            'medline': r'(?:medline|pubmed)',
            'embase': r'embase',
            'central': r'(?:central|cochrane)',
            'cinahl': r'cinahl',
            'web of science': r'web of science',
            'psycinfo': r'psycinfo'
        }
        
        # Look for search strategy sections
        strategy_sections = re.findall(
            r'(?:search strategy|database search|electronic search)[^.]*?([^.]{100,500})',
            text, re.IGNORECASE | re.MULTILINE | re.DOTALL
        )
        
        for section in strategy_sections:
            for db_name, pattern in db_patterns.items():
                if re.search(pattern, section, re.IGNORECASE):
                    # Extract date ranges
                    date_match = re.search(r'(\d{4})[^\d]*(\d{4})', section)
                    dates = f"{date_match.group(1)}–{date_match.group(2)}" if date_match else None
                    
                    # Extract search terms (simplified)
                    terms_match = re.search(r'(?:terms?|keywords?)[:\s]+([^.]{50,200})', section, re.IGNORECASE)
                    strategy = terms_match.group(1).strip() if terms_match else f"Database search for {db_name}"
                    
                    strategies.append(SearchDescriptor(
                        db=db_name.title(),
                        dates=dates,
                        strategy=strategy,
                        limits=["English"] if "english" in section.lower() else []
                    ))
        
        return strategies[:5]  # Limit to 5 databases
    
    def extract_flow_diagram(self, text: str) -> Optional[FlowCounts]:
        """Extract PRISMA flow diagram numbers using pattern matching."""
        
        # Common flow diagram patterns
        flow_patterns = {
            'identified': r'(?:identified|records found)[^\d]*(\d+)',
            'deduplicated': r'(?:after.*?duplicates?.*?removed|deduplicated)[^\d]*(\d+)',
            'screened': r'(?:screened|title.*?abstract)[^\d]*(\d+)',
            'fulltext': r'(?:full.?text|retrieved)[^\d]*(\d+)',
            'included': r'(?:included|final)[^\d]*(\d+)',
        }
        
        flow_data = {}
        for field, pattern in flow_patterns.items():
            matches = re.findall(pattern, text, re.IGNORECASE)
            if matches:
                try:
                    flow_data[field] = int(matches[0])
                except ValueError:
                    continue
        
        # Extract exclusion reasons
        excluded = []
        exclusion_patterns = [
            r'excluded[^\d]*(\d+)[^:]*:?\s*([^.\n]{10,100})',
            r'(\d+)[^:]*excluded[^:]*:?\s*([^.\n]{10,100})'
        ]
        
        for pattern in exclusion_patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            for match in matches:
                try:
                    n = int(match[0])
                    reason = match[1].strip()
                    excluded.append(ExclusionReason(reason=reason, n=n))
                except ValueError:
                    continue
        
        if flow_data:
            return FlowCounts(
                **flow_data,
                excluded=excluded[:10] if excluded else None  # Limit exclusion reasons
            )
        
        return None

class PDFProcessor:
    """Handles extraction of systematic review content from PDF documents."""
    
    def __init__(self):
        self.text_extractor = TextExtractor()
    
    def extract_manuscript(self, pdf_path: Path) -> Optional[Manuscript]:
        """
        Extract structured manuscript data from PDF.
        
        Args:
            pdf_path: Path to the systematic review PDF
            
        Returns:
            Manuscript object with extracted data, or None if extraction fails
        """
        # TODO: Implement PDF text extraction
        # For now, return None - PDF extraction requires additional libraries
        print(f"PDF extraction not yet implemented for {pdf_path}")
        return None

class WordProcessor:
    """Handles extraction from Word documents (.docx)."""
    
    def __init__(self):
        self.text_extractor = TextExtractor()
    
    def extract_manuscript(self, docx_path: Path) -> Optional[Manuscript]:
        """Extract manuscript data from Word document."""
        if not DOCX_AVAILABLE:
            print("python-docx not available. Install with: pip install python-docx")
            return None
        
        try:
            doc = Document(docx_path)
            
            # Extract all text content
            full_text = self._extract_full_text(doc)
            
            # Extract structured components
            manuscript_id = f"docx-{docx_path.stem}"
            title = self._extract_title(doc)
            pico = self.text_extractor.extract_pico_elements(full_text)
            search_strategies = self.text_extractor.parse_search_strategies(full_text)
            flow = self.text_extractor.extract_flow_diagram(full_text)
            studies = self._extract_studies_from_tables(doc)
            
            return Manuscript(
                manuscript_id=manuscript_id,
                title=title,
                question=pico,
                search=search_strategies,
                flow=flow,
                included_studies=studies
            )
            
        except Exception as e:
            print(f"Error processing Word document {docx_path}: {e}")
            return None
    
    def _extract_full_text(self, doc: DocxDocument) -> str:
        """Extract all text content from Word document."""
        text_parts = []
        
        # Extract from paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        
        # Extract from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    text_parts.append(row_text)
        
        return "\n".join(text_parts)
    
    def _extract_title(self, doc: DocxDocument) -> Optional[str]:
        """Extract document title from first paragraph or heading."""
        for paragraph in doc.paragraphs[:5]:  # Check first 5 paragraphs
            text = paragraph.text.strip()
            if len(text) > 20 and len(text) < 200:  # Reasonable title length
                return text
        return None
    
    def _extract_studies_from_tables(self, doc: DocxDocument) -> List[StudyRecord]:
        """Extract study records from tables in the document.""" 
        studies = []
        
        for table in doc.tables:
            if len(table.rows) < 2:  # Need header + data rows
                continue
            
            # Check if this looks like a results table
            header_row = table.rows[0]
            headers = [cell.text.lower().strip() for cell in header_row.cells]
            
            # Look for study characteristics table
            if any(keyword in " ".join(headers) for keyword in ['study', 'author', 'year', 'n', 'participants']):
                studies.extend(self._parse_study_characteristics_table(table))
            
            # Look for results/outcomes table  
            elif any(keyword in " ".join(headers) for keyword in ['outcome', 'result', 'effect', 'rr', 'or', 'md']):
                study_outcomes = self._parse_results_table(table)
                # Merge with existing studies or create new ones
                for study_id, outcomes in study_outcomes.items():
                    existing = next((s for s in studies if s.study_id == study_id), None)
                    if existing:
                        existing.outcomes.extend(outcomes)
                    else:
                        studies.append(StudyRecord(study_id=study_id, outcomes=outcomes))
        
        return studies[:20]  # Limit to 20 studies
    
    def _parse_study_characteristics_table(self, table: Table) -> List[StudyRecord]:
        """Parse a table containing study characteristics."""
        studies = []
        headers = [cell.text.lower().strip() for cell in table.rows[0].cells]
        
        # Map common column names
        col_mapping = {}
        for i, header in enumerate(headers):
            if any(kw in header for kw in ['study', 'author', 'first author']):
                col_mapping['study'] = i
            elif any(kw in header for kw in ['design', 'type']):
                col_mapping['design'] = i  
            elif any(kw in header for kw in ['n', 'participants', 'sample', 'size']):
                col_mapping['n'] = i
        
        for row in table.rows[1:]:  # Skip header
            if len(row.cells) <= max(col_mapping.values(), default=-1):
                continue
            
            try:
                study_id = row.cells[col_mapping.get('study', 0)].text.strip()
                if not study_id or len(study_id) > 50:  # Skip invalid entries
                    continue
                
                design = row.cells[col_mapping.get('design', 1)].text.strip() if 'design' in col_mapping else None
                
                n_total = None
                if 'n' in col_mapping:
                    n_text = row.cells[col_mapping['n']].text.strip()
                    n_match = re.search(r'(\d+)', n_text)
                    if n_match:
                        n_total = int(n_match.group(1))
                
                studies.append(StudyRecord(
                    study_id=study_id,
                    design=design,
                    n_total=n_total,
                    outcomes=[]
                ))
                
            except (ValueError, IndexError):
                continue
        
        return studies
    
    def _parse_results_table(self, table: Table) -> Dict[str, List[OutcomeEffect]]:
        """Parse a table containing results/outcomes data."""
        study_outcomes = defaultdict(list)
        headers = [cell.text.lower().strip() for cell in table.rows[0].cells]
        
        # Map outcome columns
        col_mapping = {}
        for i, header in enumerate(headers):
            if any(kw in header for kw in ['study', 'author']):
                col_mapping['study'] = i
            elif any(kw in header for kw in ['outcome', 'endpoint']):
                col_mapping['outcome'] = i
            elif any(kw in header for kw in ['effect', 'estimate', 'rr', 'or', 'md', 'smd']):
                col_mapping['effect'] = i
            elif any(kw in header for kw in ['ci', 'confidence', 'se', 'std']):
                col_mapping['variance'] = i
        
        for row in table.rows[1:]:  # Skip header
            if len(row.cells) <= max(col_mapping.values(), default=-1):
                continue
            
            try:
                study_id = row.cells[col_mapping.get('study', 0)].text.strip()
                outcome_name = row.cells[col_mapping.get('outcome', 1)].text.strip()
                
                if not study_id or not outcome_name:
                    continue
                
                # Extract effect size
                effect_text = row.cells[col_mapping.get('effect', 2)].text.strip()
                effect_match = re.search(r'([-+]?\d*\.?\d+)', effect_text)
                if not effect_match:
                    continue
                effect = float(effect_match.group(1))
                
                # Extract variance (simplified - assume from CI or SE)
                var_text = row.cells[col_mapping.get('variance', 3)].text.strip()
                variance = 0.05  # Default variance
                
                # Try to parse confidence interval
                ci_match = re.search(r'([-+]?\d*\.?\d+)[,\s]+([-+]?\d*\.?\d+)', var_text)
                if ci_match:
                    ci_low, ci_high = float(ci_match.group(1)), float(ci_match.group(2))
                    se = (ci_high - ci_low) / (2 * 1.96)  # Approximate SE from 95% CI
                    variance = se ** 2
                else:
                    # Try to parse SE directly  
                    se_match = re.search(r'([-+]?\d*\.?\d+)', var_text)
                    if se_match:
                        se = float(se_match.group(1))
                        variance = se ** 2
                
                # Determine effect metric from context
                effect_metric = "logRR"  # Default
                if any(kw in effect_text.lower() for kw in ['odds', 'or']):
                    effect_metric = "logOR"
                elif any(kw in effect_text.lower() for kw in ['mean', 'md']):
                    effect_metric = "MD"
                elif any(kw in effect_text.lower() for kw in ['std', 'smd']):
                    effect_metric = "SMD"
                
                study_outcomes[study_id].append(OutcomeEffect(
                    name=outcome_name,
                    effect_metric=effect_metric,
                    effect=effect,
                    var=variance
                ))
                
            except (ValueError, IndexError):
                continue
        
        return dict(study_outcomes)

# Factory function for document processing
def create_processor(file_path: Path):
    """Create appropriate processor based on file extension."""
    suffix = file_path.suffix.lower()
    if suffix == '.pdf':
        return PDFProcessor()
    elif suffix in ['.docx', '.doc']:
        return WordProcessor()
    else:
        raise ValueError(f"Unsupported file type: {suffix}")

# Main extraction function
def extract_manuscript_from_file(file_path: Path) -> Optional[Manuscript]:
    """Extract manuscript from file using appropriate processor."""
    processor = create_processor(file_path)
    return processor.extract_manuscript(file_path)
