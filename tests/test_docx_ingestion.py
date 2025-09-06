"""Tests for DOCX ingestion functionality."""

import pytest
from pathlib import Path
from app.utils.pdf_ingest import extract_manuscript_from_file, TextExtractor

# Sample systematic review text for testing
SAMPLE_SYSTEMATIC_REVIEW_TEXT = """
Title: Effects of ACE Inhibitors on Chronic Kidney Disease Progression: A Systematic Review

INTRODUCTION
Chronic kidney disease (CKD) affects millions worldwide and often progresses to end-stage renal disease.

METHODS

Population: Adults aged 18-75 years with chronic kidney disease stages 3-4
Intervention: ACE inhibitor therapy (enalapril, lisinopril, or ramipril)  
Comparator: Placebo or standard care without ACE inhibitors
Outcomes: Primary outcome was eGFR decline ≥40% at 12 months; Secondary outcomes included all-cause mortality and cardiovascular events

Search Strategy: We searched MEDLINE from inception to 2024, EMBASE from 1990 to 2024, and CENTRAL from inception to 2024. 
Search terms included: (ACE inhibitor OR angiotensin converting enzyme inhibitor) AND (chronic kidney disease OR CKD OR renal insufficiency) AND (progression OR decline)

PRISMA Flow: A total of 1,247 records were identified through database searching. After removing 423 duplicates, 824 records were screened by title and abstract. 156 full-text articles were retrieved for detailed evaluation. Finally, 12 studies were included in the quantitative synthesis.

RESULTS
The included studies enrolled a total of 3,456 participants. Most studies were randomized controlled trials with follow-up periods ranging from 6 to 24 months.

Table 1: Study Characteristics
Study | Design | N | Intervention | Follow-up
Smith 2019 | RCT | 234 | Enalapril 10mg | 12 months  
Johnson 2020 | RCT | 187 | Lisinopril 5mg | 18 months
Lee 2021 | RCT | 156 | Ramipril 2.5mg | 12 months

Table 2: Primary Outcomes
Study | Outcome | Effect Size | 95% CI | P-value
Smith 2019 | eGFR decline ≥40% | 0.75 | 0.58, 0.97 | 0.03
Johnson 2020 | eGFR decline ≥40% | 0.68 | 0.49, 0.94 | 0.02  
Lee 2021 | All-cause mortality | 0.82 | 0.61, 1.11 | 0.19

CONCLUSIONS
ACE inhibitors significantly reduce the risk of CKD progression in adults with stages 3-4 CKD.
"""

def test_text_extractor_pico():
    """Test PICO extraction from sample text."""
    extractor = TextExtractor()
    pico = extractor.extract_pico_elements(SAMPLE_SYSTEMATIC_REVIEW_TEXT)
    
    assert pico is not None
    assert pico.framework == "PICO"
    assert "adults aged 18-75" in pico.population.lower()
    assert "ace inhibitor" in pico.intervention.lower()
    assert "placebo" in pico.comparator.lower()
    assert len(pico.outcomes) >= 2
    assert any("egfr" in outcome.lower() for outcome in pico.outcomes)

def test_text_extractor_search_strategies():
    """Test search strategy extraction.""" 
    extractor = TextExtractor()
    strategies = extractor.parse_search_strategies(SAMPLE_SYSTEMATIC_REVIEW_TEXT)
    
    assert len(strategies) >= 2  # Should find MEDLINE, EMBASE, etc.
    db_names = [s.db.lower() for s in strategies]
    assert any("medline" in name for name in db_names)
    assert any("embase" in name for name in db_names)
    
    # Check date extraction
    dated_strategies = [s for s in strategies if s.dates]
    assert len(dated_strategies) >= 1

def test_text_extractor_flow_diagram():
    """Test PRISMA flow extraction."""
    extractor = TextExtractor()
    flow = extractor.extract_flow_diagram(SAMPLE_SYSTEMATIC_REVIEW_TEXT)
    
    assert flow is not None
    assert flow.identified == 1247
    assert flow.deduplicated == 824  # 1247 - 423 
    assert flow.screened == 824
    assert flow.fulltext == 156
    assert flow.included == 12

def test_docx_availability():
    """Test that DOCX dependencies are available for testing."""
    try:
        from docx import Document
        from app.utils.pdf_ingest import DOCX_AVAILABLE
        assert DOCX_AVAILABLE is True
    except ImportError:
        pytest.skip("python-docx not available - install with: pip install python-docx")

@pytest.mark.skipif(not Path("tests/sample_manuscript.json").exists(), reason="Sample file not found")
def test_docx_ingestion_flow():
    """Test the complete DOCX ingestion workflow (requires python-docx)."""
    
    try:
        from docx import Document
    except ImportError:
        pytest.skip("python-docx not available")
    
    # Create a minimal DOCX for testing
    doc = Document()
    doc.add_heading('Systematic Review of ACE Inhibitors', 0)
    
    # Add PICO section
    doc.add_heading('Methods', 1)
    doc.add_paragraph('Population: Adults with CKD stages 3-4')
    doc.add_paragraph('Intervention: ACE inhibitor therapy')
    doc.add_paragraph('Comparator: Placebo or standard care')  
    doc.add_paragraph('Outcomes: eGFR decline ≥40%, all-cause mortality')
    
    # Add search section
    doc.add_paragraph('Search Strategy: MEDLINE 1990-2024, EMBASE 1990-2024')
    
    # Add flow 
    doc.add_paragraph('PRISMA Flow: 500 records identified, 300 after duplicates removed, 50 full-text reviewed, 10 included')
    
    # Add simple table
    table = doc.add_table(rows=2, cols=3)
    table.cell(0, 0).text = 'Study'
    table.cell(0, 1).text = 'Design'  
    table.cell(0, 2).text = 'N'
    table.cell(1, 0).text = 'Test2024'
    table.cell(1, 1).text = 'RCT'
    table.cell(1, 2).text = '100'
    
    # Save temporary file
    test_docx_path = Path("tests/temp_test.docx")
    doc.save(test_docx_path)
    
    try:
        # Test extraction
        manuscript = extract_manuscript_from_file(test_docx_path)
        
        assert manuscript is not None
        assert manuscript.manuscript_id.startswith("docx-")
        assert manuscript.title is not None
        assert manuscript.question is not None
        assert manuscript.question.population is not None
        assert manuscript.question.intervention is not None
        
        # Check that we extracted some search strategies or flow data
        assert len(manuscript.search) > 0 or manuscript.flow is not None
        
    finally:
        # Clean up
        if test_docx_path.exists():
            test_docx_path.unlink()

def test_invalid_file_handling():
    """Test handling of non-existent files."""
    from app.utils.pdf_ingest import extract_manuscript_from_file
    
    result = extract_manuscript_from_file(Path("nonexistent.docx"))
    assert result is None

def test_table_parsing_edge_cases():
    """Test table parsing with various edge cases."""
    try:
        from docx import Document
    except ImportError:
        pytest.skip("python-docx not available")
        
    from app.utils.pdf_ingest import WordProcessor
    
    doc = Document()
    
    # Add empty table
    empty_table = doc.add_table(rows=1, cols=2)
    empty_table.cell(0, 0).text = "Header1"
    empty_table.cell(0, 1).text = "Header2"
    
    # Add table with invalid data
    invalid_table = doc.add_table(rows=2, cols=3)
    invalid_table.cell(0, 0).text = "Study"
    invalid_table.cell(0, 1).text = "Effect" 
    invalid_table.cell(0, 2).text = "CI"
    invalid_table.cell(1, 0).text = "InvalidStudy"
    invalid_table.cell(1, 1).text = "not_a_number"
    invalid_table.cell(1, 2).text = "invalid_ci"
    
    processor = WordProcessor()
    studies = processor._extract_studies_from_tables(doc)
    
    # Should handle gracefully without crashing
    assert isinstance(studies, list)
    # May be empty due to invalid data, but should not crash