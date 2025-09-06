"""
Utility to create test manuscript fixtures for testing.

This creates minimal DOCX files for testing edge cases and error handling.
"""

from docx import Document
from pathlib import Path


def create_minimal_docx():
    """Create a minimal DOCX file for testing."""
    doc = Document()
    doc.add_heading('Minimal Test Document', 0)
    doc.add_paragraph('This is a minimal test document.')
    
    fixtures_dir = Path(__file__).parent
    fixtures_dir.mkdir(exist_ok=True)
    
    output_path = fixtures_dir / "minimal_test.docx"
    doc.save(output_path)
    print(f"Created minimal test document: {output_path}")
    return output_path


def create_mock_systematic_review():
    """Create a mock systematic review document for testing."""
    doc = Document()
    
    # Title
    doc.add_heading('Effectiveness of Test Interventions: A Systematic Review and Meta-Analysis', 0)
    
    # Abstract
    doc.add_heading('Abstract', 1)
    doc.add_paragraph('Background: This is a test systematic review.')
    doc.add_paragraph('Methods: We searched multiple databases.')
    doc.add_paragraph('Results: We found test results.')
    doc.add_paragraph('Conclusions: Test interventions may be effective.')
    
    # PICO section
    doc.add_heading('Research Question (PICO)', 1)
    doc.add_paragraph('Population: Adult patients with test condition (age ≥18 years)')
    doc.add_paragraph('Intervention: Test intervention (drug X, device Y)')
    doc.add_paragraph('Comparator: Standard care or placebo')
    doc.add_paragraph('Outcomes: Primary outcome at 6 months, secondary outcomes at 12 months')
    
    # Methods
    doc.add_heading('Methods', 1)
    doc.add_heading('Search Strategy', 2)
    doc.add_paragraph('MEDLINE (via PubMed): (test[MeSH] OR "test condition"[tiab]) AND intervention[tiab]')
    doc.add_paragraph('Search dates: January 2010 to December 2023')
    
    # PRISMA Flow
    doc.add_heading('Study Selection', 2)
    doc.add_paragraph('Records identified through database searching (n=1,234)')
    doc.add_paragraph('Records after duplicates removed (n=987)')
    doc.add_paragraph('Records screened (n=987)')
    doc.add_paragraph('Full-text articles assessed for eligibility (n=45)')
    doc.add_paragraph('Studies included in qualitative synthesis (n=12)')
    doc.add_paragraph('Studies included in quantitative synthesis (meta-analysis) (n=8)')
    
    # Study characteristics table
    doc.add_heading('Study Characteristics', 2)
    
    table = doc.add_table(rows=1, cols=5)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Study ID'
    hdr_cells[1].text = 'Design'
    hdr_cells[2].text = 'N'
    hdr_cells[3].text = 'Intervention'
    hdr_cells[4].text = 'Outcome'
    
    # Add test studies
    studies = [
        ('Smith 2020', 'RCT', '120', 'Drug X 10mg', 'Response rate 65%'),
        ('Jones 2021', 'RCT', '200', 'Drug X 20mg', 'Response rate 72%'),
        ('Brown 2022', 'Cohort', '150', 'Device Y', 'Improvement 58%'),
        ('Wilson 2023', 'RCT', '180', 'Drug X 15mg', 'Response rate 68%'),
    ]
    
    for study_id, design, n, intervention, outcome in studies:
        row_cells = table.add_row().cells
        row_cells[0].text = study_id
        row_cells[1].text = design
        row_cells[2].text = n
        row_cells[3].text = intervention
        row_cells[4].text = outcome
    
    # Results
    doc.add_heading('Results', 1)
    doc.add_paragraph('Meta-analysis of 8 studies showed...')
    
    fixtures_dir = Path(__file__).parent
    fixtures_dir.mkdir(exist_ok=True)
    
    output_path = fixtures_dir / "mock_systematic_review.docx"
    doc.save(output_path)
    print(f"Created mock systematic review: {output_path}")
    return output_path


def create_malformed_docx():
    """Create a DOCX with intentionally malformed structure for testing."""
    doc = Document()
    
    # Missing clear structure
    doc.add_paragraph('Random text without clear systematic review structure')
    doc.add_paragraph('Population: Unclear population description')
    doc.add_paragraph('No clear intervention described')
    doc.add_paragraph('Outcomes: Vague outcomes')
    doc.add_paragraph('Some numbers: 100, 50, 25 (unclear what they represent)')
    
    # Malformed table
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = 'Study'
    table.cell(0, 1).text = 'Data'
    table.cell(1, 0).text = 'Incomplete'
    # Leave second cell empty
    
    fixtures_dir = Path(__file__).parent
    fixtures_dir.mkdir(exist_ok=True)
    
    output_path = fixtures_dir / "malformed_test.docx"
    doc.save(output_path)
    print(f"Created malformed test document: {output_path}")
    return output_path


def main():
    """Create all test fixtures."""
    print("Creating test manuscript fixtures...")
    
    create_minimal_docx()
    create_mock_systematic_review()  
    create_malformed_docx()
    
    print("All test fixtures created successfully!")


if __name__ == "__main__":
    main()